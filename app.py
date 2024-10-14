from flask import Flask, render_template, request, send_file, make_response, redirect, url_for
from docx import Document
from docx.shared import Inches
import openai
from io import BytesIO


app = Flask(__name__)

def askGPT(text, doc):
    openai.api_key = "YOUR API KEY"
    response = openai.Completion.create(
        engine = "text-davinci-003",
        prompt = text,
        temperature = 0.6,
        max_tokens = 150
    )
    answer = response.choices[0].text.strip()
    doc.add_heading("Question:", level=1)
    doc.add_paragraph(text)
    doc.add_heading("Answer:", level=1)
    doc.add_paragraph(answer)
    return answer
    

@app.route('/')
def home():
    return render_template('index.html')

@app.route('/get_answer', methods=['POST'])
def get_answer():
    myQn = request.form['question']
    doc = Document()
    answer = askGPT(myQn, doc)
    #doc.save("questions_answers.docx")
    return {'answer': answer}
	

@app.route('/download_answers', methods=['POST'])
def download_answers():
    chatlogs_html = request.form['chatlogs']
    
    # Create a new Word document
    document = Document()
    
    # Add a title to the document
    document.add_heading('Chat Logs', 0)
    
    # Parse the HTML and find all the chat messages
    chat_messages = chatlogs_html.split('<div class="chat ')
    chat_messages = [msg for msg in chat_messages if '<p class="chat-message">' in msg]
    
    # Loop through the chat messages and add them to the document
    for i in range(2, len(chat_messages), 2):
        # Get the question and answer
        question = chat_messages[i-1].split('</p>')[0].replace('<p class="chat-message">', '')
        answer = chat_messages[i].split('</p>')[0].replace('<p class="chat-message">', '')

        # Removing Extra HTML tags 
        question = question.replace('user"><div class="user-photo"></div>','')
        answer = answer.replace('friend"><div class="user-photo"></div>','')
		
        # Add the question and answer to the document
        document.add_heading('Question', level=1)
        document.add_paragraph(question)
        document.add_heading('Answer', level=1)
        document.add_paragraph(answer)
    
    # Save the document
    #document.save("name.docx")
    
	# Save the updated document to a BytesIO object
    file_stream = BytesIO()
    document.save(file_stream)

    # Set the BytesIO object's position to the beginning
    file_stream.seek(0)

    # Create a Flask response object with the sanitized document and download options
    response = make_response(file_stream.getvalue())
    response.headers['Content-Disposition'] = 'attachment; filename=sanitized.docx'
    response.headers['Content-Type'] = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'

    return response

if __name__ == '__main__':
    app.run(debug=True)